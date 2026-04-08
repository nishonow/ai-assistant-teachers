import asyncio
import logging
import os
import re
import uuid
import aiofiles
from sqlalchemy import delete, insert, select
from app.models import Document, Chunk
from app.services.chunking import split_text
from app.services.embeddings import embed_chunks
from app.database import SessionLocal

UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 20 * 1024 * 1024
EMBEDDING_BATCH_SIZE = 32

logger = logging.getLogger(__name__)


async def save_file(file) -> tuple[str, str]:
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise ValueError("File too large. Max 20MB.")
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    ext = file.filename.rsplit(".", 1)[-1].lower()
    safe_name = f"{uuid.uuid4().hex}.{ext}"
    file_path = os.path.join(UPLOAD_DIR, safe_name)
    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)
    return file.filename, file_path


def clean_text(text: str) -> str:
    text = re.sub(r'cbd\.minjust\.gov\.kg', '', text)
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def extract_segments(file_path: str, file_type: str) -> list[tuple[str, int | None]]:
    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            cleaned = clean_text(f.read())
        return [(cleaned, None)] if cleaned else []

    elif file_type == "pdf":
        import pdfplumber
        segments: list[tuple[str, int | None]] = []
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                page_text = clean_text(page.extract_text() or "")
                if page_text:
                    segments.append((page_text, page_number))
        return segments

    elif file_type == "docx":
        import docx
        doc = docx.Document(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        cleaned = clean_text("\n\n".join(parts))
        return [(cleaned, None)] if cleaned else []

    else:
        raise ValueError(f"Unsupported file type: {file_type}")


async def _store_embedded_batch(
    *,
    session,
    document_id: int,
    chunk_items: list[tuple[str, int | None]],
    start_index: int,
) -> int:
    embeddings = await embed_chunks([chunk_text for chunk_text, _ in chunk_items])
    if len(embeddings) != len(chunk_items):
        raise RuntimeError("Embedding count does not match extracted chunk count")

    rows = [
        {
            "document_id": document_id,
            "chunk_index": start_index + offset,
            "page_number": page_number,
            "chunk_text": chunk_text,
            "embedding": embedding,
        }
        for offset, ((chunk_text, page_number), embedding) in enumerate(zip(chunk_items, embeddings))
    ]
    await session.execute(insert(Chunk), rows)
    await session.commit()
    return len(rows)


async def process_document(document_id: int):
    async with SessionLocal() as session:
        try:
            document = (
                await session.execute(select(Document).where(Document.id == document_id))
            ).scalar_one_or_none()
            if not document:
                return

            document.status = "processing"
            await session.commit()

            # Defensive reset in case a previous indexing run was interrupted.
            await session.execute(delete(Chunk).where(Chunk.document_id == document.id))
            await session.commit()

            segments = await asyncio.to_thread(extract_segments, document.file_path, document.file_type)

            pending_batch: list[tuple[str, int | None]] = []
            indexed_chunk_count = 0
            next_chunk_index = 0
            for segment_text, page_number in segments:
                for chunk_text in split_text(segment_text):
                    normalized_chunk = chunk_text.strip()
                    if normalized_chunk:
                        pending_batch.append((normalized_chunk, page_number))
                    if len(pending_batch) < EMBEDDING_BATCH_SIZE:
                        continue

                    inserted_count = await _store_embedded_batch(
                        session=session,
                        document_id=document.id,
                        chunk_items=pending_batch,
                        start_index=next_chunk_index,
                    )
                    indexed_chunk_count += inserted_count
                    next_chunk_index += inserted_count
                    pending_batch.clear()

            if pending_batch:
                inserted_count = await _store_embedded_batch(
                    session=session,
                    document_id=document.id,
                    chunk_items=pending_batch,
                    start_index=next_chunk_index,
                )
                indexed_chunk_count += inserted_count
                next_chunk_index += inserted_count
                pending_batch.clear()

            if indexed_chunk_count == 0:
                raise ValueError("No extractable text found in document")

            document.status = "indexed"
            await session.commit()
            logger.info("Document indexed successfully: document_id=%s chunks=%s", document_id, indexed_chunk_count)

        except Exception:
            logger.exception("Document processing failed for document_id=%s", document_id)
            await session.rollback()
            try:
                document = (
                    await session.execute(select(Document).where(Document.id == document_id))
                ).scalar_one_or_none()
                if document:
                    await session.execute(delete(Chunk).where(Chunk.document_id == document_id))
                    document.status = "failed"
                    await session.commit()
            except Exception:
                await session.rollback()
